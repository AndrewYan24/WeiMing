import OCR
import zhconv
from docx import Document

# 继承变量
directory = OCR.directory
pdf_page = OCR.pdf_page
pdf_name = OCR.pdf_name
rel_output_path = OCR.rel_output_path

# 创建docx文件
doc = Document()
doc.save(rel_output_path + '/' + pdf_name + '.docx')
print('开始输出docx文件')
title = doc.add_heading(pdf_name, 1)
title.add_run('\n - 微明工具生成')
now_page = 1
for doc_ot in range(pdf_page):  # 按PDF页数循环输出OCR结果到docx文件
    now_page_content = str(directory[now_page])
    # 识别句读时替换各种奇怪符号为句号
    now_page_content = now_page_content.replace('o', '。')
    now_page_content = now_page_content.replace('◎', '。')
    now_page_content = now_page_content.replace('0', '。')
    # 转换繁简体
    now_page_content = zhconv.convert(now_page_content, 'zh-cn')
    if now_page == 1:
        para = doc.add_paragraph(now_page_content)
        para.add_run('\n')
    else:
        para.add_run(now_page_content)
        para.add_run('\n')
    now_page = now_page + 1
    doc.save(rel_output_path + '/' + pdf_name + '.docx')
print('转换完毕')
